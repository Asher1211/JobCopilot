import io
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from docx import Document

DOCX_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


async def parse_resume(file_bytes: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()

    if suffix == ".docx":
        return _parse_docx(file_bytes)
    elif suffix == ".pdf":
        return _parse_pdf(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {suffix}")


def _parse_docx(file_bytes: bytes) -> str:
    # Method 1: Standard python-docx extraction
    doc = Document(io.BytesIO(file_bytes))
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if para.style.name.startswith("Heading"):
                level = para.style.name.split()[-1]
                paragraphs.append(f"{'#' * int(level)} {text}")
            else:
                paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    result = "\n\n".join(paragraphs)
    if len(result.strip()) >= 100:
        return result

    # Method 2: Raw XML extraction (handles WPS Office, text boxes, special formatting)
    with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
        xml_bytes = z.read("word/document.xml")

    root = ET.fromstring(xml_bytes)
    all_text: list[str] = []
    for elem in root.iter(f"{{{DOCX_NS}}}t"):
        t = (elem.text or "").strip()
        if t:
            all_text.append(t)

    return "\n".join(all_text)


def _parse_pdf(file_bytes: bytes) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages: list[str] = []

    for page in doc:
        text = page.get_text("text")
        if text.strip():
            pages.append(text.strip())

    doc.close()
    return "\n\n".join(pages)
